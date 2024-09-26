import pandas as pd
from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from docx import Document
from docx.shared import Pt

# Inicializa a janela principal
janela = Tk()
janela.title("Gerador de Certificado")

# Configuração do estilo
stilo = ttk.Style()
stilo.theme_use("alt")
stilo.configure(".", font=("Segoe UI", 12), rowheight=30)
janela.configure(bg="#e7f1ff")  # Fundo azul claro

# Criação da Treeview para exibir dados
treeviewDados = ttk.Treeview(janela, columns=(1, 2, 3, 4, 5, 6), show="headings")
treeviewDados.column("1", anchor=CENTER)
treeviewDados.heading("1", text="CPF")
treeviewDados.column("2", anchor=CENTER)
treeviewDados.heading("2", text="Nome")
treeviewDados.column("3", anchor=CENTER)
treeviewDados.heading("3", text="RG")
treeviewDados.column("4", anchor=CENTER)
treeviewDados.heading("4", text="Data Inicio")
treeviewDados.column("5", anchor=CENTER)
treeviewDados.heading("5", text="Data Fim")
treeviewDados.column("6", anchor=CENTER)
treeviewDados.heading("6", text="Email")

treeviewDados.grid(row=4, column=0, columnspan=6, sticky="NSEW", pady=15)

# Função para passar dados da Treeview para os campos de entrada
def funcaoPassaDadosTreeviewParaEntry(Event):
    item = treeviewDados.selection()
    for i in item:
        exibirCPF.delete(0, END)
        exibirNome.delete(0, END)
        exibirRG.delete(0, END)
        exibirDataInicio.delete(0, END)
        exibirDataFim.delete(0, END)
        exibirEmail.delete(0, END)

        exibirCPF.insert(0, treeviewDados.item(i, "values")[0])
        exibirNome.insert(0, treeviewDados.item(i, "values")[1])
        exibirRG.insert(0, treeviewDados.item(i, "values")[2])
        exibirDataInicio.insert(0, treeviewDados.item(i, "values")[3])
        exibirDataFim.insert(0, treeviewDados.item(i, "values")[4])
        exibirEmail.insert(0, treeviewDados.item(i, "values")[5])

treeviewDados.bind("<Double-1>", funcaoPassaDadosTreeviewParaEntry)

# Abrindo o arquivo de dados
dadosUsuarios = pd.read_excel("C:\\Users\\User\\Downloads\\Dados.xlsx")

# Convertendo as colunas de data para texto
dadosUsuarios["Data Inicio"] = dadosUsuarios["Data Inicio"].astype(str)
dadosUsuarios["Data Fim"] = dadosUsuarios["Data Fim"].astype(str)

# Preenchendo a Treeview com os dados do Excel
for linha in range(len(dadosUsuarios)):
    dataInicioTratada = dadosUsuarios.iloc[linha, 3].split("-")
    dataInicioTratada = f"{dataInicioTratada[2]}/{dataInicioTratada[1]}/{dataInicioTratada[0]}"
    
    dataFimTratada = dadosUsuarios.iloc[linha, 4].split("-")
    dataFimTratada = f"{dataFimTratada[2]}/{dataFimTratada[1]}/{dataFimTratada[0]}"
    
    treeviewDados.insert("", "end",
                         values=(str(dadosUsuarios.iloc[linha, 0]),  # CPF
                                 str(dadosUsuarios.iloc[linha, 1]),  # Nome
                                 str(dadosUsuarios.iloc[linha, 2]),  # RG
                                 str(dataInicioTratada),  # Data Inicio
                                 str(dataFimTratada),  # Data Fim
                                 str(dadosUsuarios.iloc[linha, 5])))  # Email

# Campos de entrada
Label(janela, text="CPF:", font=("Segoe UI", 12), bg="#e7f1ff").grid(row=0, column=0, sticky="E", pady=10)
exibirCPF = Entry(janela, font=("Segoe UI", 12))
exibirCPF.grid(row=0, column=1, sticky="W", pady=10)

Label(janela, text="Nome:", font=("Segoe UI", 12), bg="#e7f1ff").grid(row=0, column=2, sticky="E", pady=10)
exibirNome = Entry(janela, font=("Segoe UI", 12))
exibirNome.grid(row=0, column=3, sticky="W", pady=10)

Label(janela, text="RG:", font=("Segoe UI", 12), bg="#e7f1ff").grid(row=0, column=4, sticky="E", pady=10)
exibirRG = Entry(janela, font=("Segoe UI", 12))
exibirRG.grid(row=0, column=5, sticky="W", pady=10)

Label(janela, text="Data Inicio:", font=("Segoe UI", 12), bg="#e7f1ff").grid(row=1, column=0, sticky="E", pady=10)
exibirDataInicio = Entry(janela, font=("Segoe UI", 12))
exibirDataInicio.grid(row=1, column=1, sticky="W", pady=10)

Label(janela, text="Data Fim:", font=("Segoe UI", 12), bg="#e7f1ff").grid(row=1, column=2, sticky="E", pady=10)
exibirDataFim = Entry(janela, font=("Segoe UI", 12))
exibirDataFim.grid(row=1, column=3, sticky="W", pady=10)

Label(janela, text="Email:", font=("Segoe UI", 12), bg="#e7f1ff").grid(row=1, column=4, sticky="E", pady=10)
exibirEmail = Entry(janela, font=("Segoe UI", 12))
exibirEmail.grid(row=1, column=5, sticky="W", pady=10)

# Função para filtrar dados na Treeview
def filtrarDados():
    for linha in treeviewDados.get_children():
        treeviewDados.delete(linha)

    if exibirCPF.get() == "":
        for linha in range(len(dadosUsuarios)):
            dataInicioTratada = dadosUsuarios.iloc[linha, 3].split("-")
            dataInicioTratada = f"{dataInicioTratada[2]}/{dataInicioTratada[1]}/{dataInicioTratada[0]}"

            dataFimTratada = dadosUsuarios.iloc[linha, 4].split("-")
            dataFimTratada = f"{dataFimTratada[2]}/{dataFimTratada[1]}/{dataFimTratada[0]}"

            treeviewDados.insert("", "end",
                                 values=(str(dadosUsuarios.iloc[linha, 0]),
                                         str(dadosUsuarios.iloc[linha, 1]),
                                         str(dadosUsuarios.iloc[linha, 2]),
                                         str(dataInicioTratada),
                                         str(dataFimTratada),
                                         str(dadosUsuarios.iloc[linha, 5])))
    else:
        for linha in range(len(dadosUsuarios)):
            if exibirCPF.get() == str(dadosUsuarios.iloc[linha, 0]):
                dataInicioTratada = dadosUsuarios.iloc[linha, 3].split("-")
                dataInicioTratada = f"{dataInicioTratada[2]}/{dataInicioTratada[1]}/{dataInicioTratada[0]}"

                dataFimTratada = dadosUsuarios.iloc[linha, 4].split("-")
                dataFimTratada = f"{dataFimTratada[2]}/{dataFimTratada[1]}/{dataFimTratada[0]}"

                treeviewDados.insert("", "end",
                                     values=(str(dadosUsuarios.iloc[linha, 0]),
                                             str(dadosUsuarios.iloc[linha, 1]),
                                             str(dadosUsuarios.iloc[linha, 2]),
                                             str(dataInicioTratada),
                                             str(dataFimTratada),
                                             str(dadosUsuarios.iloc[linha, 5])))

# Botões para pesquisa e geração de certificados
botaoPesquisar = Button(janela, text="PESQUISAR", font=("Segoe UI", 12), command=filtrarDados, bg="#6bbf9a", fg="white")
botaoPesquisar.grid(row=5, column=0, sticky="NSEW", padx=10, pady=10)

def gerarCertificado():
    arquivoWord = Document("C:\\Users\\User\\Downloads\\Certificado.docx")

    # Pegando os dados do aluno
    nomeAluno = exibirNome.get()
    dataInicio = exibirDataInicio.get()
    dataFim = exibirDataFim.get()
    CPF_Aluno = exibirCPF.get()
    RG_Aluno = exibirRG.get()

    frase_parte1 = " concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de "
    frase_montada = f"{nomeAluno}, CPF: {CPF_Aluno}, RG: {RG_Aluno}, {frase_parte1} {dataInicio} a {dataFim}."

    estilo = arquivoWord.styles["Normal"]

    for paragrafo in arquivoWord.paragraphs:
        if "@nome" in paragrafo.text:
            paragrafo.text = paragrafo.text.replace("@nome", nomeAluno)
            fonte = estilo.font
            fonte.name = "Calibri"
            fonte.size = Pt(24)
        if "@DataInicio" in paragrafo.text:
            paragrafo.text = paragrafo.text.replace("@DataInicio", dataInicio)
            fonte = estilo.font
            fonte.name = "Calibri"
            fonte.size = Pt(24)
        if "@DataFim" in paragrafo.text:
            paragrafo.text = paragrafo.text.replace("@DataFim", dataFim)
            fonte = estilo.font
            fonte.name = "Calibri"
            fonte.size = Pt(24)
        if "@CPF" in paragrafo.text:
            paragrafo.text = paragrafo.text.replace("@CPF", CPF_Aluno)
            fonte = estilo.font
            fonte.name = "Calibri"
            fonte.size = Pt(24)
        if "@RG" in paragrafo.text:
            paragrafo.text = paragrafo.text.replace("@RG", RG_Aluno)
            fonte = estilo.font
            fonte.name = "Calibri"
            fonte.size = Pt(24)

    caminhoCertificadoGerado = f"C:\\Users\\User\\Downloads\\certificado_{nomeAluno.replace(' ', '_')}.docx"
    arquivoWord.save(caminhoCertificadoGerado)
    messagebox.showinfo("Mensagem", "Certificado gerado com sucesso!")

# Função para gerar certificados em massa
def gerarCertificadosEmMassa():
    arquivoWordModelo = "C:\\Users\\User\\Downloads\\Certificado.docx"
    for i in range(len(dadosUsuarios)):
        nomeAluno_Separado = str(dadosUsuarios.iloc[i, 1])
        CPF_Aluno = str(dadosUsuarios.iloc[i, 0])
        RG_Aluno = str(dadosUsuarios.iloc[i, 2])
        dataInicio = str(dadosUsuarios.iloc[i, 3]).split("-")
        dataInicio = f"{dataInicio[2]}/{dataInicio[1]}/{dataInicio[0]}"
        dataFim = str(dadosUsuarios.iloc[i, 4]).split("-")
        dataFim = f"{dataFim[2]}/{dataFim[1]}/{dataFim[0]}"

        # Gerar o certificado
        arquivoWord = Document(arquivoWordModelo)
        estilo = arquivoWord.styles["Normal"]

        for paragrafo in arquivoWord.paragraphs:
            if "@nome" in paragrafo.text:
                paragrafo.text = paragrafo.text.replace("@nome", nomeAluno_Separado)
                fonte = estilo.font
                fonte.name = "Calibri"
                fonte.size = Pt(24)
            if "@DataInicio" in paragrafo.text:
                paragrafo.text = paragrafo.text.replace("@DataInicio", dataInicio)
                fonte = estilo.font
                fonte.name = "Calibri"
                fonte.size = Pt(24)
            if "@DataFim" in paragrafo.text:
                paragrafo.text = paragrafo.text.replace("@DataFim", dataFim)
                fonte = estilo.font
                fonte.name = "Calibri"
                fonte.size = Pt(24)
            if "@CPF" in paragrafo.text:
                paragrafo.text = paragrafo.text.replace("@CPF", CPF_Aluno)
                fonte = estilo.font
                fonte.name = "Calibri"
                fonte.size = Pt(24)
            if "@RG" in paragrafo.text:
                paragrafo.text = paragrafo.text.replace("@RG", RG_Aluno)
                fonte = estilo.font
                fonte.name = "Calibri"
                fonte.size = Pt(24)

        caminhoCertificadoGerado = f"C:\\Users\\User\\Downloads\\certificado_{nomeAluno_Separado.replace(' ', '_')}.docx"
        arquivoWord.save(caminhoCertificadoGerado)

    # Mensagem de sucesso após gerar todos os certificados
    messagebox.showinfo("Mensagem", "Certificados gerados em massa com sucesso!")

# Botões para gerar certificados
botaoGerarEmMassa = Button(janela, text="Gerar em Massa", font=("Segoe UI", 12), command=gerarCertificadosEmMassa, bg="#6bbf9a", fg="white")
botaoGerarEmMassa.grid(row=6, column=0, sticky="NSEW", padx=10, pady=10)

botaoGerarCertificado = Button(janela, text="Gerar Certificado", font=("Segoe UI", 12), command=gerarCertificado, bg="#6bbf9a", fg="white")
botaoGerarCertificado.grid(row=6, column=1, sticky="NSEW", padx=10, pady=10)

# Ajustando o layout da janela
janela.grid_rowconfigure(7, weight=1)
janela.grid_columnconfigure(0, weight=1)
janela.grid_columnconfigure(1, weight=1)

# Inicia a aplicação
janela.mainloop()
